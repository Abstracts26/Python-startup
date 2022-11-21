from docx import Document
from docx.shared import Inches

document = Document()

#information:
name = input('What is your name? ')
ph_no= input ('What is your phone number? ')
email= input ('What is your email address? ') 
document.add_picture ('me.png', width=Inches (2))
document.add_paragraph (name + ' | ' + ph_no + ' | ' + email)

#About me
document.add_heading('About Me')
B= input ('Tell something about yourself- ')
document.add_paragraph (B)

#Education:
document.add_heading('Education')
p= document.add_paragraph()
Educational_experience= input ('Enter your educational experience- ')
Course= input('Enter the course name- ')
Institute= input('Enter Institute name- ')
Passing_year= input('Enter Passing Year- ')
p.add_run (Educational_experience).bold= True
p.add_run (' in '+Course+ '\n    ').bold= True
p.add_run (Institute+'                       ')== True
CGPA= input ('Your GPA at '+ Educational_experience + '- ')
p.add_run (CGPA + '\n    ').italic=True
p.add_run (Passing_year).italic= True



#More Education:

while True:
          Education= input ("Do you have more education to add? - YES OR NO: ")

          if Education.lower() == 'yes':
                              
                            p= document.add_paragraph ()
                            Educational_experience= input ('Enter your educational experience- ')
                            Course= input('Enter the course name- ')
                            Institute= input('Enter Institute name- ')
                            Passing_year= input('Enter Passing Year- ')
                            p.add_run (Educational_experience).bold= True
                            p.add_run ('\n    '+Course+ '\n    ')== True
                            p.add_run (Institute+'                       ')== True
                            CGPA= input ('Your GPA at '+ Educational_experience + '- ')
                            p.add_run (CGPA + '\n    ').italic=True
                            p.add_run (Passing_year).italic= True

          else:
               break

#Work Experience:
document.add_heading ('Work Experiences')
p= document.add_paragraph ()
   
Experience_details= input('Enter the designation- ')
From_date= input('From- ')
To_date= input('To- ')

p.add_run (Experience_details+ ('                       ')).bold= True
p.add_run (From_date+ (' - ')+To_date + ('\n')).italic= True

Company= input ('The company you were at- ')
p.add_run (Company).italic=True


# more work experiences:
while True:
    More_experiences= input ("Do you have more experiences? - YES OR NO: ")

    if More_experiences.lower() == 'yes':
                        
                        p= document.add_paragraph ()

                        Experience_details= input('Enter the designation- ')
                        From_date= input('From- ')
                        To_date= input('To- ')

                        p.add_run (Experience_details+ ('                       ')).bold= True
                        p.add_run (From_date+ (' - ')+To_date + ('\n')).italic= True

                        Company= input ('The company you were at- ')
                        p.add_run (Company).italic=True

    else:
        break


#Skills:

document.add_heading ('Skills')
skill= input ('Enter your skills- ')
p= document.add_paragraph (skill)
p.style='List Bullet'

# more skills:
while True:
    More_skills= input ("Do you have more skills? - YES OR NO: ")

    if More_skills.lower() == 'yes':
                        
                        skill= input ('Enter your skills- ')
                        p= document.add_paragraph (skill)
                        p.style='List Bullet'

    else:
        break                    
    
    
document.save ('Cv.docx')